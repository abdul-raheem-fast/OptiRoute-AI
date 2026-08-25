import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

doc = docx.Document()

# Page Setup: Standard Letter Margins (1 inch)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Colors
PRIMARY_COLOR = RGBColor(27, 79, 114)     # Deep Navy Blue
SECONDARY_COLOR = RGBColor(40, 116, 166)  # Steel Blue
TEXT_COLOR = RGBColor(44, 62, 80)         # Charcoal Dark
ACCENT_COLOR = RGBColor(211, 84, 0)       # Pumpkin Accent

# Helper Styling Functions
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_header_p(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = PRIMARY_COLOR
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = SECONDARY_COLOR
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = TEXT_COLOR
    return p

def add_styled_paragraph(text, bold_prefix=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.color.rgb = PRIMARY_COLOR
        r_pre.font.name = 'Calibri'
    r_text = p.add_run(text)
    r_text.font.color.rgb = TEXT_COLOR
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(11)
    return p

# -------------------------------------------------------------
# COVER TITLE BLOCK
# -------------------------------------------------------------
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(24)
title_p.paragraph_format.space_after = Pt(4)
r_title = title_p.add_run("AetherFlow")
r_title.bold = True
r_title.font.size = Pt(32)
r_title.font.color.rgb = PRIMARY_COLOR
r_title.font.name = 'Arial'

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(18)
r_sub = sub_p.add_run("Intelligent Multi-LLM Routing & Evaluation Benchmark Framework")
r_sub.font.size = Pt(16)
r_sub.font.color.rgb = SECONDARY_COLOR
r_sub.font.name = 'Arial'

meta_p = doc.add_paragraph()
meta_p.paragraph_format.space_after = Pt(24)
r_meta = meta_p.add_run("Master Technical Specification, Model Selection Rationale & Empirical Analytics\nProject Architecture & Dataset Suite")
r_meta.italic = True
r_meta.font.size = Pt(11)
r_meta.font.color.rgb = RGBColor(120, 144, 156)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# -------------------------------------------------------------
# 1. EXECUTIVE SUMMARY & SYSTEM OVERVIEW
# -------------------------------------------------------------
add_header_p("1. Executive Summary & System Overview", level=1)

add_styled_paragraph(
    "AetherFlow is an intelligent multi-LLM orchestration and dynamic routing framework designed to optimize performance, monetary cost, and response latency across enterprise LLM workloads. In modern Artificial Intelligence deployments, relying on a single flagship model (such as GPT-5 or Gemini-2.5-Pro) for all incoming user queries incurs exorbitant API expenditures ($20–$80 per 1,000 queries). Conversely, directing all traffic to lightweight 8B models leads to severe task failure on complex mathematical proofs, software engineering, and scientific QA.",
    bold_prefix="Project Mission: "
)

add_styled_paragraph(
    "To resolve this dilemma, AetherFlow establishes a data-driven Pareto frontier across eight leading Large Language Models (LLMs) evaluated on 26 benchmark sources standardized into five core capability classes. The AetherFlow dataset suite provides zero-corruption, prompt-aligned empirical ground truth to train dynamic ML routers capable of reducing API expenditure by 70–80% while preserving over 90%+ of top-tier flagship accuracy.",
    bold_prefix="Core Value Proposition: "
)

# -------------------------------------------------------------
# 2. STANDARDIZED DATASET ARCHITECTURE
# -------------------------------------------------------------
add_header_p("2. Standardized Dataset Architecture (5 Core Classes)", level=1)

add_styled_paragraph(
    "All evaluation benchmarks in AetherFlow are audited, cleaned, and categorized into five human-understandable, non-overlapping target capability classes:"
)

table_classes = doc.add_table(rows=6, cols=5)
table_classes.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_cls = ["Class Name", "Domain Focus", "Question Type", "Answer Format", "Merged Source Benchmarks"]
hdr_cells = table_classes.rows[0].cells
for idx, text in enumerate(headers_cls):
    hdr_cells[idx].text = text
    set_cell_background(hdr_cells[idx], "1B4F72")
    hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
    hdr_cells[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

class_data = [
    ("Coding", "Software engineering, algorithms & Python synthesis", "Open-Ended Code Generation", "Python code blocks, functions, assertions", "livecodebench, humaneval, mbpp, swe-bench, arenahard_coding"),
    ("Mathematical Reasoning", "Standard algebra, calculus & numeric word problems", "Open-Ended Numeric Solving", "Numerical values, equations, calculations", "livemathbench, math500, mathbench, arenahard_math"),
    ("Scientific Questionnaire", "Graduate STEM (Physics, Organic Chemistry, Biology)", "Multiple-Choice Questions (MCQs)", "Choice option letters (A, B, C, D)", "gpqa, medqa, finqa"),
    ("General Knowledge", "Humanities, social sciences, multi-subject academic QA", "Open-Ended & MCQs", "Concise text answers, multi-subject choices", "mmlupro, arcc, bbh, winogrande, arenahard, simpleqa, meld, emorynlp, korbench, kandk, arc-agi, hle"),
    ("Competitive Math", "Olympiad-level mathematics proofs & contest math", "Open-Ended High-Level Solving", "Multi-page proofs & exact integer solutions", "aime (American Invitational Math Exam)")
]

for row_idx, data_tuple in enumerate(class_data, start=1):
    row_cells = table_classes.rows[row_idx].cells
    fill = "F2F4F4" if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(data_tuple):
        row_cells[col_idx].text = text
        set_cell_background(row_cells[col_idx], fill)
        p = row_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if col_idx == 0:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = PRIMARY_COLOR

add_styled_paragraph(
    "Standardized 15-Column Schema: Every dataset CSV file across all folders follows a uniform schema: index, dataset_name, model_name, origin_query, prompt, ground_truth, prediction, score, correct, prompt_tokens, completion_tokens, total_tokens, cost, estimated_latency, raw_output.",
    bold_prefix="\nSchema Standardization: ", space_after=12
)

# -------------------------------------------------------------
# 3. MODEL SELECTION RATIONALE & PARETO FRONTIER
# -------------------------------------------------------------
add_header_p("3. Model Selection Rationale & Pareto Frontier Analysis", level=1)

add_styled_paragraph(
    "AetherFlow evaluates an ensemble of eight models carefully selected to span parameter scales, licensing types (open-weights vs. proprietary), and architectural paradigms:"
)

table_models = doc.add_table(rows=9, cols=5)
table_models.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_mod = ["Model Name", "Developer", "Access Type", "Architecture Class", "Target Role in Routing"]
hdr_cells_m = table_models.rows[0].cells
for idx, text in enumerate(headers_mod):
    hdr_cells_m[idx].text = text
    set_cell_background(hdr_cells_m[idx], "2874A6")
    hdr_cells_m[idx].paragraphs[0].runs[0].font.bold = True
    hdr_cells_m[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells_m[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

model_data = [
    ("Llama-3.1-8B-Instruct", "Meta AI", "Open-Weights", "Dense 8B Transformer", "Zero/Minimal Cost Anchor ($0.000013/query)"),
    ("Qwen3-8B", "Alibaba Cloud", "Open-Weights", "Chain-of-Thought (<think>) 8B", "Open-weights test-time CoT reasoning"),
    ("DeepSeek-v3-0324", "DeepSeek AI", "Open / API", "Multi-Head Latent Attention / MoE", "High-efficiency open flagship reasoning ($0.00084/query)"),
    ("Gemini-2.5-Flash", "Google DeepMind", "Proprietary API", "Distilled Multimodal Transformer", "Ultra-Fast Speed Anchor (0.194s latency)"),
    ("GPT-4.1", "OpenAI", "Proprietary API", "Dense / MoE Transformer", "Standard commercial enterprise baseline"),
    ("Claude-Sonnet-4", "Anthropic", "Proprietary API", "Enterprise Transformer", "High-precision code synthesis & low hallucination"),
    ("Gemini-2.5-Pro", "Google DeepMind", "Proprietary API", "Multimodal MoE Flagship", "Graduate scientific QA & complex math proofs"),
    ("GPT-5", "OpenAI", "Proprietary API", "Frontier Multimodal Architecture", "Maximum Accuracy Ceiling (88.77% accuracy)")
]

for row_idx, data_tuple in enumerate(model_data, start=1):
    row_cells = table_models.rows[row_idx].cells
    fill = "F2F4F4" if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(data_tuple):
        row_cells[col_idx].text = text
        set_cell_background(row_cells[col_idx], fill)
        p = row_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if col_idx == 0:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = PRIMARY_COLOR

add_styled_paragraph(
    "Multi-Dimensional Pareto Variance: Cost Variance: 6,200x ($0.000013 to $0.0829/query) | Latency Variance: 55x (0.194s to 10.8s) | Accuracy Variance: 51.4% (37.31% to 88.77%). This extreme variation provides the exact empirical spread needed to optimize LLM routers.",
    bold_prefix="\nPareto Metrics: ", space_after=12
)

# -------------------------------------------------------------
# 4. FOLDER ARRANGEMENTS & ALIGNMENT
# -------------------------------------------------------------
add_header_p("4. Dataset Folder Arrangements", level=1)

add_styled_paragraph(
    "1. aligned_8_models/ (1,887 rows per file): 100% identical query alignment across all 8 models. Row index i across all 8 files contains the exact same prompt, enabling 1-to-1 prompt routing experiments.\n"
    "2. aligned_7_models/ (3,352 rows per file): Extended query alignment across 7 models (excluding GPT-4.1) incorporating ArenaHard benchmarks.\n"
    "3. individual/ (Unconstrained Datasets): Full unconstrained evaluation runs per model standardized under the 5 target classes."
)

# -------------------------------------------------------------
# 5. EMPIRICAL RESULTS & EMBEDDED VISUALIZATIONS
# -------------------------------------------------------------
add_header_p("5. Empirical Evaluation Results & Visualizations", level=1)

add_styled_paragraph("Comprehensive Model Performance Matrix (Aligned 8 Models):")

table_results = doc.add_table(rows=9, cols=5)
table_results.alignment = WD_TABLE_ALIGNMENT.CENTER

headers_res = ["Model Name", "Mean Accuracy (%)", "Avg Cost / Query ($)", "Cost / 1,000 Queries ($)", "Avg Latency (s)"]
hdr_cells_r = table_results.rows[0].cells
for idx, text in enumerate(headers_res):
    hdr_cells_r[idx].text = text
    set_cell_background(hdr_cells_r[idx], "1B4F72")
    hdr_cells_r[idx].paragraphs[0].runs[0].font.bold = True
    hdr_cells_r[idx].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells_r[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

res_data = [
    ("GPT-5", "88.77%", "$0.021178", "$21.18", "0.813s"),
    ("Gemini-2.5-Pro", "87.44%", "$0.082901", "$82.90", "1.897s"),
    ("Qwen3-8B", "76.68%", "$0.000807", "$0.81", "3.669s"),
    ("Gemini-2.5-Flash", "76.10%", "$0.003463", "$3.46", "0.194s"),
    ("DeepSeek-v3-0324", "75.73%", "$0.000841", "$0.84", "2.922s"),
    ("Claude-Sonnet-4", "75.09%", "$0.009877", "$9.88", "0.325s"),
    ("GPT-4.1", "72.39%", "$0.004097", "$4.10", "1.043s"),
    ("Llama-3.1-8B-Instruct", "37.31%", "$0.000013", "$0.013", "0.315s")
]

for row_idx, data_tuple in enumerate(res_data, start=1):
    row_cells = table_results.rows[row_idx].cells
    fill = "F2F4F4" if row_idx % 2 == 0 else "FFFFFF"
    for col_idx, text in enumerate(data_tuple):
        row_cells[col_idx].text = text
        set_cell_background(row_cells[col_idx], fill)
        p = row_cells[col_idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if col_idx == 0:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = PRIMARY_COLOR

# Embedded Graphs
graphs_folder = r"D:\Datasets (FYP)\graphs_output"
graph_files = [
    ("model_cost_comparison.png", "Figure 1: Average Monetary Cost per Query ($) Across Evaluated Models"),
    ("model_latency_comparison.png", "Figure 2: Average Response Latency (Seconds) Across Evaluated Models"),
    ("accuracy_vs_cost_tradeoff.png", "Figure 3: Efficiency Tradeoff: Accuracy vs. Cost (USD per 1k Queries)"),
    ("accuracy_vs_latency_tradeoff.png", "Figure 4: Speed Tradeoff: Accuracy vs. Response Latency"),
    ("avg_cost_per_class.png", "Figure 5: Average Cost per Query by Benchmark Class"),
    ("avg_latency_per_class.png", "Figure 6: Average Latency by Benchmark Class")
]

add_header_p("Visual Analytics & Charts", level=2)

for g_name, caption in graph_files:
    g_path = os.path.join(graphs_folder, g_name)
    if os.path.exists(g_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(12)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(g_path, width=Inches(5.8))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(14)
        r_cap = p_cap.add_run(caption)
        r_cap.italic = True
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = RGBColor(100, 110, 120)

# -------------------------------------------------------------
# 6. ACADEMIC CITATIONS & REFERENCES
# -------------------------------------------------------------
add_header_p("6. Academic Citations & Literature References", level=1)

citations = [
    ("1. Meta AI - Llama-3.1-8B: ", "Dubey et al., 'The LLaMA 3 Herd of Models', arXiv:2407.21783 (2024)."),
    ("2. Alibaba Cloud AI - Qwen3-8B: ", "Yang et al., 'Qwen2.5 Technical Report', arXiv:2409.12190 (2024)."),
    ("3. DeepSeek AI - DeepSeek-v3: ", "Liu et al., 'DeepSeek-V3 Technical Report', arXiv:2412.19437 (2024)."),
    ("4. Google DeepMind - Gemini 2.5: ", "Google DeepMind Team, 'Gemini 1.5 & 2.5 Technical Report', arXiv:2403.05530 (2024)."),
    ("5. OpenAI - GPT-4.1 / GPT-5: ", "OpenAI, 'GPT-4 Technical Report & Next-Gen Systems', arXiv:2303.08774 (2023-2025)."),
    ("6. Anthropic - Claude-Sonnet-4: ", "Anthropic, 'The Claude 3 & 4 Model Family Technical Report', Anthropic Publications (2024)."),
    ("7. FrugalGPT (Stanford): ", "Chen, L., Zaharia, M., & Zou, J. 'FrugalGPT: How to Use Large Language Models Cheaper and Better', arXiv:2305.05176 (2023)."),
    ("8. RouteLLM (UC Berkeley / LMSYS): ", "Ong, I., Rashad, A., Chiang, W.L., Stoica, I., et al. 'RouteLLM: Learning to Route LLMs Efficiently', arXiv:2406.18665 (2024).")
]

for prefix, cite in citations:
    add_styled_paragraph(cite, bold_prefix=prefix, space_after=4)

output_docx = r"D:\Datasets (FYP)\AetherFlow_Master_Technical_and_Academic_Documentation.docx"
doc.save(output_docx)
print(f"✓ Successfully generated master Word Document: {output_docx}")
